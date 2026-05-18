"""
Build submittal-flow templates matching the actual Centric / EMA visual style.

Style rules pulled from the live submittal PDFs:

  Transmittal
    - Left-aligned letterhead (small), Centric URL underlined in hyperlink blue
    - Centered black bold "TRANSMITTAL" title (plain, no banner)
    - Bordered tables. Label cells filled light blue (#DAE3F3), black text
    - At the bottom: red A/E REVIEW STAMP block (filled in by Engineer)

  Cover Sheet / Index
    - Centered black bold title (no banner)
    - Subtitle "Division ## ## ##  -  [Name]" in blue (#1F3864)
    - Thin blue horizontal rule under the title block
    - Bold black ALL-CAPS section headers, each underlined with a thin
      gray bottom border
    - KV pairs as plain bold-label / plain-value text (no fill-in line)
    - Index also includes Contractor's Remarks + an italic footer line
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.normpath(os.path.join(HERE, "..", "templates", "submittals"))
os.makedirs(OUT, exist_ok=True)

# Palette pulled from the real PDFs
LABEL_BLUE = "DAE3F3"
SUBTITLE_BLUE = RGBColor(0x1F, 0x38, 0x64)
RULE_BLUE = "1F3864"
RULE_GRAY = "BFBFBF"
STAMP_RED = RGBColor(0xC0, 0x00, 0x00)
STAMP_RED_HEX = "C00000"
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HYPER_BLUE = RGBColor(0x05, 0x63, 0xC1)


# -----------------------------------------------------------------------------
# XML / formatting helpers
# -----------------------------------------------------------------------------

def shd(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill_hex)
    tcPr.append(s)


def cell_border(cell, edges, color_hex, sz="6"):
    """Set explicit borders on a single cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is None:
        existing = OxmlElement('w:tcBorders')
        tcPr.append(existing)
    for edge in edges:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), color_hex)
        existing.append(b)


def table_borders(table, color_hex="7F7F7F", sz="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color_hex)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_run(run, bold=False, italic=False, size=10, color=BLACK,
            name="Calibri", underline=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, **kw):
    align = kw.pop('align', None)
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), **kw)
    return p


def set_col_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = Inches(widths[idx])


def cell_text(cell, text, bold=False, italic=False, size=10, color=BLACK,
              fill=None, align=None, underline=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), bold=bold, italic=italic, size=size,
            color=color, underline=underline)
    if fill is not None:
        shd(cell, fill)


def hr_blue(doc):
    """Thin blue horizontal rule (Cover Sheet / Index title separator)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), RULE_BLUE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_heading(doc, text):
    """Bold black uppercase heading with a thin gray bottom border."""
    p = doc.add_paragraph()
    set_run(p.add_run(text.upper()), bold=True, size=11)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), RULE_GRAY)
    pBdr.append(bottom)
    pPr.append(pBdr)


def kv_plain(doc, rows, label_width=1.6, value_width=5.4):
    """Plain bold-label / plain-value KV table - no borders, no fill."""
    t = doc.add_table(rows=len(rows), cols=2)
    for i, (k, v) in enumerate(rows):
        cell_text(t.rows[i].cells[0], k, bold=True, size=10)
        cell_text(t.rows[i].cells[1], v, size=10)
    set_col_widths(t, [label_width, value_width])
    return t


# -----------------------------------------------------------------------------
# Letterhead
# -----------------------------------------------------------------------------

def letterhead(doc):
    add_para(doc, "[Company Name]", bold=True, size=11)
    add_para(doc, "[Street Address]", size=9)
    add_para(doc, "[City, State ZIP]", size=9)
    add_para(doc, "[Phone]", size=9)
    p = doc.add_paragraph()
    set_run(p.add_run("[Website]"), size=9, color=HYPER_BLUE, underline=True)


# -----------------------------------------------------------------------------
# Cover Sheet / Index title block (centered black title + blue subtitle + rule)
# -----------------------------------------------------------------------------

def title_block(doc, title_text, subtitle_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(title_text), bold=True, size=18)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p2.add_run(subtitle_text), size=11, color=SUBTITLE_BLUE)
    hr_blue(doc)


# -----------------------------------------------------------------------------
# TRANSMITTAL
# -----------------------------------------------------------------------------

def section_transmittal(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("TRANSMITTAL"), bold=True, size=20)
    doc.add_paragraph()

    # DATE row
    t = doc.add_table(rows=1, cols=2)
    table_borders(t)
    cell_text(t.rows[0].cells[0], "DATE:", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[0].cells[1], "[YYYY-MM-DD]", size=10)
    set_col_widths(t, [1.0, 6.1])
    doc.add_paragraph()

    # TO / VIA / PROJECT block (6 rows x 4 cols)
    t = doc.add_table(rows=6, cols=4)
    table_borders(t)
    cell_text(t.rows[0].cells[0], "TO:", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[0].cells[1], "[Recipient Firm]", size=10)
    cell_text(t.rows[0].cells[2], "VIA:", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[0].cells[3], "U.S. MAIL", size=10)
    cell_text(t.rows[1].cells[0], "ATTN:", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[1].cells[1], "[Recipient Name]", size=10)
    cell_text(t.rows[1].cells[2], "", size=10)
    cell_text(t.rows[1].cells[3], "HAND DELIVERED", size=10)
    cell_text(t.rows[2].cells[0], "", size=10)
    cell_text(t.rows[2].cells[1], "[Street Address]", size=10)
    cell_text(t.rows[2].cells[2], "", size=10)
    cell_text(t.rows[2].cells[3], "LOCAL DELIVERY", size=10)
    cell_text(t.rows[3].cells[0], "", size=10)
    cell_text(t.rows[3].cells[1], "[City, State ZIP]", size=10)
    cell_text(t.rows[3].cells[2], "", size=10)
    cell_text(t.rows[3].cells[3], "OVERNIGHT SHIPMENT", size=10)
    cell_text(t.rows[4].cells[0], "", size=10)
    cell_text(t.rows[4].cells[1], "", size=10)
    cell_text(t.rows[4].cells[2], "x", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[4].cells[3], "Email", size=10)
    cell_text(t.rows[5].cells[0], "PROJECT:", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[5].cells[1], "[Project Name]", size=10)
    cell_text(t.rows[5].cells[2], "PROJECT #", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[5].cells[3], "[Project #]", size=10)
    set_col_widths(t, [0.9, 3.0, 1.0, 2.2])
    doc.add_paragraph()

    # TRANSMITTED / ITEM grid
    t = doc.add_table(rows=4, cols=5)
    table_borders(t)
    cell_text(t.rows[0].cells[0], "TRANSMITTED:", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[0].cells[1], "x", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[0].cells[2], "HEREWITH", size=10)
    cell_text(t.rows[0].cells[3], "", size=10)
    cell_text(t.rows[0].cells[4], "SEPARATELY", size=10)
    cell_text(t.rows[1].cells[0], "ITEM:", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[1].cells[1], "", size=10)
    cell_text(t.rows[1].cells[2], "ORIGINAL DRAWINGS", size=10)
    cell_text(t.rows[1].cells[3], "", size=10)
    cell_text(t.rows[1].cells[4], "COPY OF LETTER", size=10)
    cell_text(t.rows[2].cells[0], "", size=10)
    cell_text(t.rows[2].cells[1], "", size=10)
    cell_text(t.rows[2].cells[2], "PRINTS", size=10)
    cell_text(t.rows[2].cells[3], "", size=10)
    cell_text(t.rows[2].cells[4], "SPECIFICATIONS", size=10)
    cell_text(t.rows[3].cells[0], "", size=10)
    cell_text(t.rows[3].cells[1], "x", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[3].cells[2], "SUBMITTALS", size=10)
    cell_text(t.rows[3].cells[3], "", size=10)
    cell_text(t.rows[3].cells[4], "SAMPLES", size=10)
    set_col_widths(t, [1.2, 0.4, 2.1, 0.4, 3.0])
    doc.add_paragraph()

    # QUANTITY / DATED / DESCRIPTION
    t = doc.add_table(rows=5, cols=3)
    table_borders(t)
    cell_text(t.rows[0].cells[0], "QUANTITY", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[0].cells[1], "DATED", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[0].cells[2], "DESCRIPTION", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[1].cells[0], "[#]", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[1].cells[1], "[YYYY-MM-DD]", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[1].cells[2],
              "[Description of the transmitted item / submittal package]", size=10)
    for r in (2, 3, 4):
        for c in range(3):
            cell_text(t.rows[r].cells[c], "", size=10)
    set_col_widths(t, [1.0, 1.2, 4.9])
    doc.add_paragraph()

    # REMARKS
    t = doc.add_table(rows=1, cols=2)
    table_borders(t)
    cell_text(t.rows[0].cells[0], "REMARKS:", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[0].cells[1],
              "[One paragraph summarizing what is enclosed, the spec section "
              "it satisfies, and any context the reviewer needs (revision "
              "number, scope coverage, related RFIs, requested action).]",
              size=10)
    set_col_widths(t, [1.2, 5.9])
    doc.add_paragraph()

    # By line
    t = doc.add_table(rows=1, cols=2)
    table_borders(t)
    cell_text(t.rows[0].cells[0], "By:", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[0].cells[1], "[Name, Title]", size=10)
    set_col_widths(t, [1.0, 6.1])
    doc.add_paragraph()
    doc.add_paragraph()

    # A/E REVIEW STAMP (red box, filled in by the Engineer)
    add_review_stamp(doc)


def add_review_stamp(doc):
    """Red review stamp area filled in by the Architect/Engineer on return."""
    # 4-row stamp: checkbox grid (2 rows), disclaimer paragraph (1 row), sig (1 row)
    t = doc.add_table(rows=4, cols=2)
    for row in t.rows:
        for cell in row.cells:
            cell_border(cell, ('top', 'left', 'bottom', 'right'),
                        STAMP_RED_HEX, sz="12")

    # Row 0: REVIEWED  |  REVISE AND RESUBMIT
    cell_text(t.rows[0].cells[0], "[ ]  REVIEWED",
              bold=True, size=10, color=STAMP_RED)
    cell_text(t.rows[0].cells[1], "[ ]  REVISE AND RESUBMIT",
              bold=True, size=10, color=STAMP_RED)

    # Row 1: FURNISH AS CORRECTED  |  REJECTED
    cell_text(t.rows[1].cells[0], "[ ]  FURNISH AS CORRECTED",
              bold=True, size=10, color=STAMP_RED)
    cell_text(t.rows[1].cells[1], "[ ]  REJECTED",
              bold=True, size=10, color=STAMP_RED)

    # Row 2: merge both cells, disclaimer paragraph
    a = t.rows[2].cells[0]
    b = t.rows[2].cells[1]
    a.merge(b)
    cell_text(a,
              "Checking is only for general conformance with the design "
              "concept of the project and general compliance with the "
              "information given in the Contract Documents. Any action shown "
              "is subject to the requirements of the plans and "
              "specifications. Contractor is responsible for dimensions "
              "which shall be confirmed and correlated at the job site; "
              "fabrication processes and techniques of construction; "
              "coordination of his work with that of all other trades; and "
              "the satisfactory performance of his work.",
              size=9, color=STAMP_RED)

    # Row 3: Date / By / A/E firm block
    cell_text(t.rows[3].cells[0],
              "Date:  ______________________\n\nBy:    ______________________",
              size=10, color=STAMP_RED)
    cell_text(t.rows[3].cells[1],
              "[A/E Firm Name]\n[City  -  City  -  City]\n[Phone]\n[Firm "
              "Registration No.]\n[Website]",
              size=8, color=STAMP_RED, align=WD_ALIGN_PARAGRAPH.RIGHT)

    set_col_widths(t, [3.5, 3.6])

    # Caption under the stamp
    doc.add_paragraph()
    add_para(doc,
             "Red review stamp area is completed by the Architect/Engineer "
             "on return - left blank by the Contractor at issue.",
             italic=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)


# -----------------------------------------------------------------------------
# COVER SHEET
# -----------------------------------------------------------------------------

def section_cover_sheet(doc):
    title_block(doc, "SUBMITTAL COVER SHEET",
                "Division [## ## ##]  -  [Division Name]")
    doc.add_paragraph()

    section_heading(doc, "Project")
    kv_plain(doc, [
        ["Project Name:", "[Project Name]"],
        ["Project / CSP Number:", "[Contract # / CSP #]"],
        ["A&E Project No.:", "[A/E Project #]"],
        ["Location:", "[Site Address]"],
        ["Date:", "[Date]"],
        ["Submittal No.:", "[Spec section - sequential, e.g. 26 00 00-01]"],
        ["Submittal Type:", "[Product Data / Shop Drawings / Samples / Engineering / Mock-up / Other]"],
    ])
    doc.add_paragraph()

    section_heading(doc, "Engineer")
    kv_plain(doc, [
        ["Firm:", "[A/E Firm]"],
        ["Attention:", "[Reviewer Name]"],
        ["Address:", "[Street, City, State ZIP]"],
    ])
    doc.add_paragraph()

    section_heading(doc, "Prime Contractor / General Contractor")
    kv_plain(doc, [
        ["Firm:", "[Prime Contractor]"],
        ["Contact:", "[Name, Title]"],
        ["Address:", "[Street, City, State ZIP]"],
        ["Phone:", "[Phone]"],
    ])
    doc.add_paragraph()

    section_heading(doc, "Subcontractor / Equipment Supplier")
    kv_plain(doc, [
        ["Firm:", "[Sub or Supplier]"],
        ["Contact:", "[Name]"],
        ["Email:", "[Email]"],
        ["PO Number:", "[PO #]"],
    ])
    doc.add_paragraph()

    section_heading(doc, "Contractor Action")
    p = doc.add_paragraph()
    set_run(p.add_run("APPROVED          APPROVED AS NOTED          REJECTED"),
            bold=True, size=11)
    doc.add_paragraph()
    add_para(doc,
             "Contractor has reviewed submitted product data against "
             "specifications. Product data forwarded to Engineer for review "
             "and comment.", size=10)
    doc.add_paragraph()

    # Signature line - underlined italic value
    t = doc.add_table(rows=1, cols=3)
    cell_text(t.rows[0].cells[0], "Contractor Signature:", bold=True, size=10)
    cell_text(t.rows[0].cells[1], "[Name, Title]",
              italic=True, size=10, underline=True)
    cell_text(t.rows[0].cells[2], "Date:  [Date]", bold=True, size=10)
    set_col_widths(t, [1.9, 2.6, 2.6])


# -----------------------------------------------------------------------------
# INDEX
# -----------------------------------------------------------------------------

def section_index(doc):
    title_block(doc, "SUBMITTAL INDEX",
                "Division [## ## ##]  -  [Division Name]")
    doc.add_paragraph()

    section_heading(doc, "Items Submitted")
    t = doc.add_table(rows=5, cols=3)
    table_borders(t)
    cell_text(t.rows[0].cells[0], "Item", bold=True, size=10, fill=LABEL_BLUE,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[0].cells[1], "Description", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[0].cells[2], "Spec Section", bold=True, size=10, fill=LABEL_BLUE,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[1].cells[0], "1", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[1].cells[1],
              "[Product / shop drawing / sample description - manufacturer, "
              "model, rating, application, quantity]", size=10)
    cell_text(t.rows[1].cells[2], "[## ## ##]", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in (2, 3, 4):
        for c in range(3):
            cell_text(t.rows[r].cells[c], "", size=10)
    set_col_widths(t, [0.7, 5.4, 1.0])
    doc.add_paragraph()

    section_heading(doc, "Contractor's Remarks")
    add_para(doc, "Division [## ## ##]  -  [Division Name]",
             size=10, color=SUBTITLE_BLUE)
    add_para(doc,
             "Contractor has reviewed submitted product data against Division "
             "[##] specifications. The following remarks are provided for "
             "the Engineer's consideration.",
             size=10)
    doc.add_paragraph()
    t = doc.add_table(rows=4, cols=4)
    table_borders(t)
    cell_text(t.rows[0].cells[0], "#", bold=True, size=10, fill=LABEL_BLUE,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[0].cells[1], "Remark", bold=True, size=10, fill=LABEL_BLUE)
    cell_text(t.rows[0].cells[2], "Spec Ref.", bold=True, size=10, fill=LABEL_BLUE,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[0].cells[3], "Status", bold=True, size=10, fill=LABEL_BLUE,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[1].cells[0], "1", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[1].cells[1],
              "[Clarification, quantity verification, or question for the "
              "Engineer]", size=10)
    cell_text(t.rows[1].cells[2], "[## ## ##]", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[1].cells[3], "FOR REVIEW", bold=True, size=9,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[2].cells[0], "2", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[2].cells[1],
              "[Item not in this submittal but covered elsewhere or available "
              "on request]", size=10)
    cell_text(t.rows[2].cells[2], "[## ## ##]", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(t.rows[2].cells[3], "FOR INFO", bold=True, size=9,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    for c in range(4):
        cell_text(t.rows[3].cells[c], "", size=10)
    set_col_widths(t, [0.5, 4.6, 1.0, 1.0])
    doc.add_paragraph()

    # Footer line
    add_para(doc,
             "Submittal No.: [## ## ##-##]  |  Date: [Date]  |  CSP #[##]  "
             "|  A/E Project #[##]",
             italic=True, size=9, color=SUBTITLE_BLUE)


# -----------------------------------------------------------------------------
# Documents
# -----------------------------------------------------------------------------

def page_setup(doc):
    s = doc.sections[0]
    s.left_margin = Inches(0.7)
    s.right_margin = Inches(0.7)
    s.top_margin = Inches(0.6)
    s.bottom_margin = Inches(0.6)


def build_transmittal():
    doc = Document()
    page_setup(doc)
    letterhead(doc)
    doc.add_paragraph()
    section_transmittal(doc)
    out = os.path.join(OUT, "01_Transmittal.docx")
    doc.save(out)
    return out


def build_cover_sheet():
    doc = Document()
    page_setup(doc)
    letterhead(doc)
    doc.add_paragraph()
    section_cover_sheet(doc)
    out = os.path.join(OUT, "02_Submittal_Cover_Sheet.docx")
    doc.save(out)
    return out


def build_index():
    doc = Document()
    page_setup(doc)
    letterhead(doc)
    doc.add_paragraph()
    section_index(doc)
    out = os.path.join(OUT, "03_Submittal_Index.docx")
    doc.save(out)
    return out


def build_combined():
    doc = Document()
    page_setup(doc)
    letterhead(doc)
    doc.add_paragraph()
    section_transmittal(doc)
    doc.add_page_break()
    section_cover_sheet(doc)
    doc.add_page_break()
    section_index(doc)
    out = os.path.join(OUT, "04_Submittal_Package_Combined.docx")
    doc.save(out)
    return out


def main():
    for fn in [build_transmittal, build_cover_sheet, build_index, build_combined]:
        print(f"wrote {os.path.basename(fn())}")


if __name__ == '__main__':
    main()

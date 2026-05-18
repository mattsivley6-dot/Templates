"""
Build clean, reusable PMBOK-aligned project artifact templates.

Output: /templates/*.docx

Each template follows PMI/PMBOK 7th Edition artifact structures and is
stripped of project-specific data (no client names, dates, or identifiers).
Placeholder text in [brackets] marks fields to fill per project.

Artifacts produced:
  01_Risk_Register.docx
  02_RFI_Log.docx
  03_Daily_Status_Report.docx
  04_Change_Request_Form.docx
  05_Project_Closeout_Checklist.docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.normpath(os.path.join(HERE, "..", "templates"))
os.makedirs(OUT, exist_ok=True)

PMI_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
PMI_ACCENT = RGBColor(0x2E, 0x75, 0xB6)
HEADER_FILL = "1F3A5F"


def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_run(run, bold=False, size=10, color=None, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run(r, bold=True, size=18, color=PMI_BLUE)


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=True, size=12, color=PMI_BLUE)


def add_para(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, bold=bold, size=size)


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    set_run(r, size=10)


def set_col_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = Inches(widths[idx])


def build_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run(r, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, HEADER_FILL)
    for ri, rowdata in enumerate(rows, start=1):
        for ci, val in enumerate(rowdata):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val) if val is not None else "")
            set_run(r, size=9)
    if widths:
        set_col_widths(t, widths)
    return t


def kv_table(doc, rows, key_width=1.8, value_width=5.4):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Light List Accent 1'
    for i, (k, v) in enumerate(rows):
        c0 = t.rows[i].cells[0]
        c1 = t.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(k)
        set_run(r0, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(c0, HEADER_FILL)
        r1 = c1.paragraphs[0].add_run(v)
        set_run(r1, size=10)
    set_col_widths(t, [key_width, value_width])
    return t


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3A5F')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_landscape(doc, left=0.5, right=0.5, top=0.5, bottom=0.5):
    section = doc.sections[0]
    new_w = section.page_height
    new_h = section.page_width
    section.page_width = new_w
    section.page_height = new_h
    section.orientation = 1
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)


def doc_control(doc, doc_type):
    add_h2(doc, "Document Control")
    kv_table(doc, [
        ["Project", "[Project Name]"],
        ["Document Type", doc_type],
        ["Document Version", "1.0"],
        ["Issue Date", "[YYYY-MM-DD]"],
        ["Prepared By", "[Name, Title, Organization]"],
        ["Distribution", "[Stakeholders / Roles]"],
        ["PMI Reference", "PMBOK Guide 7th Edition - Project Artifact"],
    ])
    doc.add_paragraph()


# -----------------------------------------------------------------------------
# 1. Risk Register
# -----------------------------------------------------------------------------

def build_risk_register():
    doc = Document()
    set_landscape(doc)
    add_title(doc, "Project Risk Register")
    add_para(doc, "PMBOK Guide 7th Edition - Uncertainty Performance Domain")
    add_horizontal_rule(doc)
    doc_control(doc, "Risk Register (PMBOK Artifact)")

    add_h2(doc, "Scoring Legend")
    add_para(doc, "Probability and Impact rated Low / Medium / High. "
                  "Risk Score = Probability x Impact (LL=1, LM/ML=2, LH/HL/MM=3, "
                  "MH/HM=4, HH=5). Strategy per PMBOK: Avoid, Transfer, Mitigate, "
                  "Accept, Escalate (threats); Exploit, Share, Enhance, Accept "
                  "(opportunities).", size=9)
    doc.add_paragraph()

    add_h2(doc, "Active Risks")
    headers = ["Risk ID", "Category", "Risk Description (Cause - Event - Effect)",
               "Risk Owner", "Prob.", "Imp.", "Score",
               "Response Strategy", "Planned Response / Contingency",
               "Trigger / Warning Sign", "Residual Risk", "Status",
               "Date Identified", "Date Reviewed", "Notes"]
    rows = [
        ["R-001", "[Category]", "[Cause] leads to [event] resulting in [effect]",
         "[Owner]", "[L/M/H]", "[L/M/H]", "",
         "[Avoid/Transfer/Mitigate/Accept/Escalate]",
         "[Planned response actions and contingency]",
         "[Trigger condition that signals the risk is materializing]",
         "[Residual risk after response]",
         "Open", "[YYYY-MM-DD]", "[YYYY-MM-DD]", ""],
    ]
    for n in range(2, 11):
        rows.append([f"R-{n:03d}", "", "", "", "", "", "", "", "", "", "",
                     "Open", "", "", ""])
    widths = [0.55, 1.1, 2.2, 0.9, 0.4, 0.4, 0.45,
              0.95, 2.4, 1.4, 1.2, 0.55, 0.85, 0.85, 1.0]
    build_table(doc, headers, rows, widths=widths)
    doc.add_paragraph()

    add_h2(doc, "Closed / Archived Risks")
    closed_rows = [["", "", "", "", "", "", "", "", "", "", "",
                    "Closed", "", "", ""]]
    build_table(doc, headers, closed_rows, widths=widths)
    doc.add_paragraph()

    add_h2(doc, "Risk Review Cadence")
    add_bullet(doc, "Daily - informal scan during team huddle")
    add_bullet(doc, "Weekly - formal review; update Status, Date Reviewed, Residual Risk")
    add_bullet(doc, "On trigger - immediate response per Planned Response column; escalate if residual remains High")

    add_h2(doc, "PMI Artifact Notes")
    add_para(doc, "Aligns with PMBOK Guide 7th Edition Risk Register and the "
                  "Uncertainty Performance Domain. Required fields: ID, Description "
                  "(cause-event-effect), Category, Probability, Impact, Score, "
                  "Strategy, Response, Trigger, Residual, Owner, Status, and "
                  "review dates.", size=9)

    out = os.path.join(OUT, "01_Risk_Register.docx")
    doc.save(out)
    return out


# -----------------------------------------------------------------------------
# 2. RFI Log
# -----------------------------------------------------------------------------

def build_rfi_log():
    doc = Document()
    set_landscape(doc)
    add_title(doc, "RFI Log - Request for Information Register")
    add_para(doc, "PMBOK Guide 7th Edition - Issue Log / RFI Log Artifact")
    add_horizontal_rule(doc)
    doc_control(doc, "RFI Log (PMBOK Issue Log derivative)")

    add_h2(doc, "Field Definitions")
    add_bullet(doc, "RFI # - sequential identifier (RFI-001, RFI-002, ...)")
    add_bullet(doc, "Priority - Low / Medium / High / Urgent (per PMI Issue Log convention)")
    add_bullet(doc, "Status - Open / In Review / Answered / Closed / Void")
    add_bullet(doc, "Required Response Date - target per spec or contract clause")
    add_bullet(doc, "Days Open - calculated; > 14 days triggers escalation")
    doc.add_paragraph()

    add_h2(doc, "RFI Register")
    headers = ["RFI #", "Date Submitted", "Submitted By", "Recipient",
               "Category", "Priority", "Subject",
               "Question / Description", "Attachments",
               "Required Response Date", "Date Answered",
               "Response / Resolution",
               "Days Open", "Cost Impact", "Schedule Impact",
               "Status", "Notes"]
    rows = []
    for n in range(1, 11):
        rows.append([f"RFI-{n:03d}", "", "", "", "", "", "", "", "", "", "",
                     "", "", "", "", "Open", ""])
    widths = [0.55, 0.8, 1.1, 0.9, 0.7, 0.55, 1.2, 2.5, 0.7,
              0.85, 0.8, 1.8, 0.55, 0.65, 0.75, 0.55, 1.0]
    build_table(doc, headers, rows, widths=widths)
    doc.add_paragraph()

    add_h2(doc, "Workflow")
    add_bullet(doc, "1. PM logs RFI on day of issuance; assigns RFI # and Priority")
    add_bullet(doc, "2. RFI transmitted to responder with return-receipt; logged in correspondence folder")
    add_bullet(doc, "3. Acknowledgement expected within 2 business days; final response per spec turnaround")
    add_bullet(doc, "4. On receipt: update Date Answered, Response, Cost/Schedule Impact, Status")
    add_bullet(doc, "5. RFIs with cost or schedule impact converted to Change Request")

    add_h2(doc, "PMI Artifact Notes")
    add_para(doc, "Uses the PMBOK Issue Log structure adapted for Requests for "
                  "Information. Required PMI fields (ID, Date Submitted, "
                  "Submitted By, Owner, Priority, Description, Required "
                  "Response, Resolution, Status) plus Cost and Schedule Impact "
                  "columns for construction contexts.", size=9)

    out = os.path.join(OUT, "02_RFI_Log.docx")
    doc.save(out)
    return out


# -----------------------------------------------------------------------------
# 3. Daily Status Report
# -----------------------------------------------------------------------------

def build_daily_report():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    add_title(doc, "Daily Project Status Report")
    add_para(doc, "PMBOK Guide 7th Edition - Project Communications / Status Report Artifact")
    add_horizontal_rule(doc)

    add_h2(doc, "Report Header")
    kv_table(doc, [
        ["Project", "[Project Name]"],
        ["Contract", "[Contract # / Type]"],
        ["Report #", "[Sequential]"],
        ["Report Date", "[YYYY-MM-DD]"],
        ["Day of Week", ""],
        ["Reporting Period", "Single workday"],
        ["Prepared By", "[Name, Role, Organization]"],
        ["Distribution", "[Stakeholders / Roles]"],
        ["Document Version", "1.0"],
    ], key_width=1.6, value_width=5.4)
    doc.add_paragraph()

    add_h2(doc, "Overall Performance Indicators (RYG)")
    add_para(doc, "Mark each as Green (on track), Yellow (at risk), or Red (off track). PMBOK Performance Domains.", size=9)
    rygs = [
        ["Schedule", "G / Y / R", ""],
        ["Cost / Budget", "G / Y / R", ""],
        ["Scope", "G / Y / R", ""],
        ["Quality", "G / Y / R", ""],
        ["Safety", "G / Y / R", ""],
        ["Stakeholder", "G / Y / R", ""],
    ]
    build_table(doc, ["Performance Domain", "Status", "Commentary"],
                rygs, widths=[1.8, 1.0, 4.4])
    doc.add_paragraph()

    add_h2(doc, "Executive Summary")
    add_para(doc, "[One paragraph - key accomplishments, issues, decisions needed today]", size=10)
    doc.add_paragraph()

    add_h2(doc, "Site Conditions / Weather")
    add_bullet(doc, "Temp high / low:")
    add_bullet(doc, "Conditions:")
    add_bullet(doc, "Wind:")
    add_bullet(doc, "Precipitation:")
    add_bullet(doc, "Work impact: None / Minor / Hold")

    add_h2(doc, "Resource Utilization - Crew Counts by Trade")
    crews = [["[Trade]", "[Sub]", "", "[Foreman]", ""] for _ in range(8)]
    build_table(doc, ["Trade", "Sub", "Workers", "Foreman", "Hours"],
                crews, widths=[1.4, 1.4, 1.0, 1.7, 0.8])
    doc.add_paragraph()

    add_h2(doc, "Work Performed This Period")
    add_bullet(doc, "[Bullet list by area / scope]")
    add_h2(doc, "Work Planned Next Period")
    add_bullet(doc, "[Planned scope for next workday]")
    add_h2(doc, "Schedule Performance")
    add_bullet(doc, "Baseline activity ID(s) / description:")
    add_bullet(doc, "Status vs. baseline: On schedule / Ahead by X / Behind by X days")
    add_bullet(doc, "Variance drivers:")
    add_bullet(doc, "Recovery plan (if behind):")
    add_h2(doc, "Cost Performance")
    add_bullet(doc, "PCO / variance activity today:")
    add_bullet(doc, "Cumulative open PCO value: $")
    add_bullet(doc, "Forecast impact to baseline cost:")
    add_h2(doc, "Quality")
    add_bullet(doc, "Inspections performed:")
    add_bullet(doc, "Rework or deficient work observed:")
    add_bullet(doc, "Corrective actions:")
    add_h2(doc, "Safety")
    add_bullet(doc, "Toolbox talk topic:")
    add_bullet(doc, "JSAs reviewed:")
    add_bullet(doc, "PPE compliance: Y / N")
    add_bullet(doc, "Incidents / near misses:")
    add_bullet(doc, "Corrective actions:")
    add_h2(doc, "Risks and Issues (Top Items from Registers)")
    add_bullet(doc, "New risks identified today (add to Risk Register):")
    add_bullet(doc, "Open issues actioned today:")
    add_bullet(doc, "Items escalated:")
    add_h2(doc, "Decisions Required")
    add_bullet(doc, "Decision needed / from whom / by when:")

    add_h2(doc, "Deliveries")
    build_table(doc, ["Time", "From", "Material", "Receiver", "Stored Where"],
                [["", "", "", "", ""]], widths=[0.7, 1.4, 2.4, 1.4, 1.4])
    doc.add_paragraph()
    add_h2(doc, "Visitors to Site")
    build_table(doc, ["Time", "Name", "Company", "Purpose"],
                [["", "", "", ""]], widths=[0.7, 2.0, 2.0, 2.6])
    doc.add_paragraph()
    add_h2(doc, "Notable Communications")
    add_bullet(doc, "Calls, emails, RFIs sent or received today")
    add_h2(doc, "Attachments / Photos")
    add_bullet(doc, "Photo file refs - store in /Daily Photos/YYYY-MM-DD/")
    add_bullet(doc, "Other attachments:")

    add_h2(doc, "Sign-off")
    build_table(doc, ["Role", "Name", "Date", "Signature"], [
        ["Prepared By", "", "", ""],
        ["Reviewed By", "", "", ""],
        ["Received By (Owner Rep)", "", "", ""],
    ], widths=[1.6, 1.8, 1.2, 2.4])
    doc.add_paragraph()

    add_h2(doc, "PMI Artifact Notes")
    add_para(doc, "Aligns with PMBOK 7th Edition communications artifact. RYG "
                  "status indicators map to Performance Domains (Schedule, "
                  "Cost, Scope, Quality, Stakeholder).", size=9)

    out = os.path.join(OUT, "03_Daily_Status_Report.docx")
    doc.save(out)
    return out


# -----------------------------------------------------------------------------
# 4. Change Request Form
# -----------------------------------------------------------------------------

def build_change_request():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    add_title(doc, "Change Request Form")
    add_para(doc, "PMBOK Guide 7th Edition - Integrated Change Control Artifact")
    add_horizontal_rule(doc)

    add_h2(doc, "1. Change Request Identification")
    kv_table(doc, [
        ["CR Number", "[CR-### or PCO-###]"],
        ["Project", "[Project Name]"],
        ["Contract", "[Contract # / Type]"],
        ["Subject / Title", "[Short description]"],
        ["Date Submitted", "[YYYY-MM-DD]"],
        ["Submitted By", "[Name, Title, Organization]"],
        ["Contact", "[Email / Phone]"],
        ["To (Reviewer)", "[Architect / Engineer / Owner Rep]"],
        ["Cc", "[Distribution]"],
    ])
    doc.add_paragraph()

    add_h2(doc, "2. Change Classification")
    kv_table(doc, [
        ["Change Type", "[ ] Scope  [ ] Schedule  [ ] Cost  [ ] Quality  [ ] Resource  [ ] Other"],
        ["Priority", "[ ] Low  [ ] Medium  [ ] High  [ ] Urgent  [ ] Emergency"],
        ["Source / Trigger", "[ ] RFI Response  [ ] Owner Directive  [ ] Field Condition  [ ] Design Clarification  [ ] Other"],
        ["Related RFI / Document", "[Fill in if applicable]"],
    ])
    doc.add_paragraph()

    add_h2(doc, "3. Description of Proposed Change")
    add_para(doc, "[One paragraph description of the work covered by this CR. "
                  "Reference the RFI or directive that triggered it if applicable.]")
    doc.add_paragraph()

    add_h2(doc, "4. Reason / Justification")
    add_para(doc, "[Per plan / spec section / field condition / owner directive. "
                  "Explain why this change is required.]")
    doc.add_paragraph()

    add_h2(doc, "5. Alternatives Considered")
    add_bullet(doc, "Alternative 1: [Description / Pros / Cons]")
    add_bullet(doc, "Alternative 2: [Description / Pros / Cons]")
    add_bullet(doc, "Do-nothing impact: [Risk and consequence of taking no action]")

    add_h2(doc, "6. Impact Analysis")
    kv_table(doc, [
        ["Scope Impact", "[Describe deliverables added / removed / modified]"],
        ["Schedule Impact", "[None / X days / TBD]"],
        ["Cost Impact", "[Total this CR: $]"],
        ["Quality Impact", "[Describe quality / performance implications]"],
        ["Risk Impact", "[New risks introduced or existing risks affected]"],
        ["Resource Impact", "[Crew / equipment / sub procurement implications]"],
    ])
    doc.add_paragraph()

    add_h2(doc, "7. Cost Breakdown")
    build_table(doc, ["Line Item", "Description", "Amount"], [
        ["Scope basis", "[One-line description]", ""],
        ["Direct labor", "", "$"],
        ["Direct material", "", "$"],
        ["Subcontractor", "", "$"],
        ["Equipment", "", "$"],
        ["Subtotal direct", "", "$"],
        ["Markup / Overhead", "", "$"],
        ["Bond / Insurance", "", "$"],
        ["Total this CR", "", "$"],
    ], widths=[1.6, 3.6, 1.6])
    doc.add_paragraph()

    add_h2(doc, "8. Recommendation")
    add_para(doc, "[Recommend approval / approval with modifications / rejection, "
                  "with rationale referencing baselines and stakeholder impact.]")
    doc.add_paragraph()

    add_h2(doc, "9. Backup Documentation Attached")
    add_bullet(doc, "[Sub quote or labor estimate]")
    add_bullet(doc, "[Material pricing]")
    add_bullet(doc, "[Photos or sketches as applicable]")
    add_bullet(doc, "[Related RFI or correspondence]")

    add_h2(doc, "10. Change Control Board Disposition")
    kv_table(doc, [
        ["Disposition", "[ ] Approved  [ ] Approved with Modifications  [ ] Rejected  [ ] Deferred  [ ] Pending"],
        ["Decision Date", ""],
        ["Decision Rationale", ""],
        ["Implementation Date", ""],
        ["Updated Baselines", "Scope / Schedule / Cost (note which baselines are revised)"],
    ])
    doc.add_paragraph()

    add_h2(doc, "11. Approval Signatures")
    build_table(doc, ["Role", "Name", "Date", "Signature"], [
        ["Requestor (Contractor)", "", "", ""],
        ["Architect / Engineer", "", "", ""],
        ["Owner", "", "", ""],
    ], widths=[1.8, 2.4, 1.0, 2.0])
    doc.add_paragraph()

    add_h2(doc, "12. Action Requested")
    add_para(doc, "[Specific action and date by which response is requested.]")
    doc.add_paragraph()

    add_h2(doc, "PMI Artifact Notes")
    add_para(doc, "Follows PMBOK 7th Edition Integrated Change Control format: "
                  "identification, classification, description, justification, "
                  "alternatives, impact analysis across scope/schedule/cost/"
                  "quality/risk/resources, cost breakdown, recommendation, "
                  "backup, CCB disposition, and approval signatures.", size=9)

    out = os.path.join(OUT, "04_Change_Request_Form.docx")
    doc.save(out)
    return out


# -----------------------------------------------------------------------------
# 5. Project Closeout Checklist
# -----------------------------------------------------------------------------

def build_closeout():
    doc = Document()
    set_landscape(doc)
    add_title(doc, "Project Closeout Checklist")
    add_para(doc, "PMBOK Guide 7th Edition - Close Project or Phase Artifact")
    add_horizontal_rule(doc)
    doc_control(doc, "Closeout Checklist (PMBOK Close Project/Phase artifact)")

    add_h2(doc, "Closeout Sections")
    add_para(doc, "Items grouped by PMBOK closeout category. Each item carries "
                  "ownership, source, status, and verification reference to "
                  "support formal acceptance and archival.", size=9)
    doc.add_paragraph()

    headers = ["Item #", "PMI Section", "Deliverable / Item",
               "Responsible Party", "Source / Contact",
               "Status", "% Complete", "Due Date", "Date Completed",
               "Verified By", "Location in Project Folder", "Notes"]

    sections = [
        ("Administrative - Permits", 4),
        ("Administrative - Acceptance", 2),
        ("Administrative - Archival", 1),
        ("Financial - Lien Waivers", 2),
        ("Financial - Final Pay App", 1),
        ("Contractual - Warranties", 4),
        ("Technical - O&M Manuals", 4),
        ("Technical - Attic Stock", 3),
        ("Technical - Controls", 2),
        ("Documentation - As-Builts", 3),
        ("Quality - Test Reports", 3),
        ("Quality - Punch List", 2),
        ("Resources - Training", 3),
        ("Lessons Learned", 1),
    ]

    rows = []
    idx = 1
    for section, count in sections:
        for _ in range(count):
            rows.append([f"C-{idx:03d}", section, "[Deliverable description]",
                         "[Owner]", "[Contact]", "Not Started", "0%",
                         "[YYYY-MM-DD]", "", "", "[Folder path]", ""])
            idx += 1

    widths = [0.55, 1.7, 2.0, 1.2, 1.2, 0.8, 0.55, 0.8, 0.8, 0.8, 1.3, 1.0]
    build_table(doc, headers, rows, widths=widths)
    doc.add_paragraph()

    add_h2(doc, "Closeout Acceptance")
    kv_table(doc, [
        ["Punch List Walk Date", "[YYYY-MM-DD]"],
        ["Punch List Resolution Target", "[YYYY-MM-DD]"],
        ["Substantial Completion Target", "[YYYY-MM-DD]"],
        ["Final Acceptance / Owner Sign-off", ""],
        ["Lessons Learned Session", "Schedule within 2 weeks of substantial completion"],
        ["Archive Date", ""],
    ], key_width=2.5, value_width=7.5)
    doc.add_paragraph()

    add_h2(doc, "PMI Artifact Notes")
    add_para(doc, "Aligns with PMBOK 7th Edition Close Project or Phase "
                  "process. Items grouped by PMI closeout category: "
                  "Administrative, Financial, Contractual, Technical, "
                  "Documentation, Quality, Resources/Training, Lessons "
                  "Learned.", size=9)

    out = os.path.join(OUT, "05_Project_Closeout_Checklist.docx")
    doc.save(out)
    return out


# -----------------------------------------------------------------------------

def main():
    for fn in [build_risk_register, build_rfi_log, build_daily_report,
               build_change_request, build_closeout]:
        path = fn()
        print(f"wrote {os.path.basename(path)}")


if __name__ == '__main__':
    main()
